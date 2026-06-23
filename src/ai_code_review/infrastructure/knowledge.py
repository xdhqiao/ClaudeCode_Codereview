from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any


WORD_RE = re.compile(r"[A-Za-z_][A-Za-z0-9_.-]{1,}")
CJK_RE = re.compile(r"[\u4e00-\u9fff]+")
RULE_RE = re.compile(r"(?m)^##\s+(?:RULE\s+)?([A-Za-z0-9_.-]+)\s*$")


@dataclass(slots=True)
class StandardChunk:
    rule_id: str
    languages: set[str]
    profile: str
    text: str
    source: Path


def _tokens(text: str) -> set[str]:
    """Tokenize identifiers and Chinese text without requiring an external index."""
    result = {token.lower() for token in WORD_RE.findall(text)}
    for sequence in CJK_RE.findall(text):
        if len(sequence) == 1:
            result.add(sequence)
            continue
        result.update(sequence[index : index + 2] for index in range(len(sequence) - 1))
        if len(sequence) >= 3:
            result.update(sequence[index : index + 3] for index in range(len(sequence) - 2))
    return result


def _parse_frontmatter(text: str) -> tuple[dict[str, Any], str]:
    if not text.startswith("---\n"):
        return {}, text
    end = text.find("\n---\n", 4)
    if end < 0:
        return {}, text
    metadata: dict[str, Any] = {}
    for line in text[4:end].splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        value = value.strip()
        if value.startswith("[") and value.endswith("]"):
            metadata[key.strip()] = [
                item.strip().strip("'\"") for item in value[1:-1].split(",") if item.strip()
            ]
        else:
            metadata[key.strip()] = value.strip("'\"")
    return metadata, text[end + 5 :]


def _split_rules(body: str, source: Path, languages: set[str], profile: str) -> list[StandardChunk]:
    matches = list(RULE_RE.finditer(body))
    if not matches:
        return [
            StandardChunk(
                rule_id=f"{source.stem.upper()}-DOCUMENT",
                languages=languages,
                profile=profile,
                text=body.strip(),
                source=source,
            )
        ]
    chunks: list[StandardChunk] = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(body)
        chunks.append(
            StandardChunk(
                rule_id=match.group(1),
                languages=languages,
                profile=profile,
                text=body[match.start() : end].strip(),
                source=source,
            )
        )
    return chunks


class KnowledgeBase:
    def __init__(self, root: Path, max_chars: int, max_rules: int = 20) -> None:
        self.root = root
        self.max_chars = max_chars
        self.max_rules = max_rules
        self._chunks: list[StandardChunk] | None = None

    def reload(self) -> None:
        chunks: list[StandardChunk] = []
        if self.root.exists():
            for path in sorted(self.root.rglob("*.md")):
                text = path.read_text(encoding="utf-8", errors="replace")
                metadata, body = _parse_frontmatter(text)
                languages = {
                    str(item).lower() for item in metadata.get("languages", ["all"])
                }
                profile = str(metadata.get("profile", "default"))
                chunks.extend(_split_rules(body, path, languages, profile))
        self._chunks = chunks

    def search(self, *, language: str, profile: str, query: str) -> str:
        if self._chunks is None:
            self.reload()
        query_tokens = _tokens(query)
        ranked: list[tuple[float, StandardChunk]] = []
        for chunk in self._chunks or []:
            if chunk.profile not in {profile, "default", "all"}:
                continue
            if "all" not in chunk.languages and language.lower() not in chunk.languages:
                continue
            chunk_tokens = _tokens(f"{chunk.rule_id}\n{chunk.text}")
            overlap = len(query_tokens & chunk_tokens)
            if query_tokens and chunk_tokens:
                lexical_score = overlap / max(len(query_tokens), 1)
            else:
                lexical_score = 0.0
            language_bonus = 0.25 if language.lower() in chunk.languages else 0.0
            profile_bonus = 0.15 if chunk.profile == profile else 0.0
            generic_bonus = 0.02 if "all" in chunk.languages else 0.0
            score = lexical_score + language_bonus + profile_bonus + generic_bonus
            if overlap > 0 or language_bonus > 0 or profile_bonus > 0:
                ranked.append((score, chunk))
        ranked.sort(key=lambda item: (-item[0], item[1].rule_id))

        selected: list[str] = []
        size = 0
        for _, chunk in ranked[: self.max_rules]:
            rendered = f"[{chunk.rule_id}] 来源: {chunk.source.name}\n{chunk.text}"
            remaining = self.max_chars - size
            if remaining <= 0:
                break
            if len(rendered) > remaining:
                if not selected:
                    selected.append(rendered[:remaining])
                break
            selected.append(rendered)
            size += len(rendered)
        return "\n\n".join(selected)


def import_docx_standard(
    source: Path,
    destination_dir: Path,
    *,
    standard_id: str,
    languages: list[str],
    profile: str,
    version: str,
) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to import Word standards") from exc

    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_.-]{0,127}", standard_id):
        raise ValueError(
            "standard_id must contain only letters, numbers, dots, underscores, and hyphens"
        )
    for field_name, value in {
        "profile": profile,
        "version": version,
        **{f"language[{index}]": item for index, item in enumerate(languages)},
    }.items():
        if "\n" in value or "\r" in value:
            raise ValueError(f"{field_name} must not contain newlines")
    source = source.expanduser().resolve()
    if not source.is_file():
        raise ValueError(f"Word standard does not exist: {source}")

    document = Document(source)
    destination_dir.mkdir(parents=True, exist_ok=True)
    output = destination_dir / f"{standard_id}.md"
    lines = [
        "---",
        f"standard_id: {standard_id}",
        f"version: {version}",
        f"profile: {profile}",
        f"languages: [{', '.join(languages or ['all'])}]",
        f"source: {source.name}",
        "---",
        "",
        f"# {standard_id}",
        "",
    ]

    current_heading = "通用规则"
    sequence = 1
    section_paragraphs: list[str] = []

    def flush_section() -> None:
        nonlocal sequence, section_paragraphs
        if not section_paragraphs:
            return
        rule_id = f"{standard_id.upper()}-{sequence:04d}"
        lines.extend(
            [
                f"## {rule_id}",
                f"章节：{current_heading}",
                "",
                "\n\n".join(section_paragraphs),
                "",
            ]
        )
        sequence += 1
        section_paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name or "标题" in style_name:
            flush_section()
            current_heading = text
            lines.extend([f"# {text}", ""])
            continue
        section_paragraphs.append(text)

    flush_section()

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(cells):
                continue
            rule_id = f"{standard_id.upper()}-T{table_index:02d}-{row_index:03d}"
            lines.extend([f"## {rule_id}", "", " | ".join(cells), ""])

    output.write_text("\n".join(lines), encoding="utf-8")
    return output
